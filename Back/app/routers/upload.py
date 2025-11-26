import hashlib, io, os, json, re, zipfile
from collections import deque
from typing import Iterator, Dict, Any, List, Tuple, Set

from ..core.runtime_cfg import get_runtime_cfg
from ..core.memlog import log_mem
import subprocess

import orjson
from fastapi import APIRouter, UploadFile, File, HTTPException, Query
from fastapi.responses import Response
from pathlib import Path

from ..core.config import (
    CORPUS_JSONL,
    INDEX_JSON,
    UPLOAD_DIR,
    OCR_LANG_DEFAULT,
    OCR_WORKERS_DEFAULT,
)
from ..core.io_utils import file_lock
from ..core.logger import logger

from ..services.indexing.index_build import build_index_json
from ..services.search.index_search import clear_index_cache, get_index_cached, load_index, load_index_cached
from ..services.helpers.normalizer import normalize_for_shingles, simple_tokens
from ..services.converters.pdf_heavy import extract_and_normalize_pdf
from ..services.converters.pdf_convert import smart_pdf_to_docx
from ..services.converters.docx_utils import extract_docx_text

router = APIRouter(prefix="/api", tags=["Operations"])

# --- optional parsers ---
try:
    from docx import Document
except Exception:
    Document = None
# ------------------------


# ---------- helpers ----------

def _chunked(seq, size: int):
    """Разбивает последовательность на куски по size."""
    for i in range(0, len(seq), size):
        yield seq[i:i + size]


def _bytes_to_text_utf8(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("utf-8", "ignore")

def _build_native_index() -> None:
    """
    Строит C++-индекс через бинарь `index_builder`.

    Ожидаемый CLI (подправь под свой index_builder.cpp, если отличается):
        index_builder <corpus_jsonl> <index_dir>

    Где:
      - corpus_jsonl = CORPUS_JSONL
      - index_dir    = директория, в которой лежит index.json
    """
    index_dir = INDEX_JSON.parent
    index_dir.mkdir(parents=True, exist_ok=True)

    log_mem(
        f"[index_build] native: before index_builder, corpus={CORPUS_JSONL}, index_dir={index_dir}"
    )
    logger.info(
        "[index_build] NATIVE start: corpus_jsonl=%s, index_dir=%s",
        CORPUS_JSONL,
        index_dir,
    )

    try:
        subprocess.run(
            [
                "index_builder",
                str(CORPUS_JSONL),
                str(index_dir),
            ],
            check=True,
        )
    except FileNotFoundError as e:
        logger.error("[index_build] native index_builder not found in PATH: %s", e)
        raise HTTPException(500, "native index_builder binary not found")
    except subprocess.CalledProcessError as e:
        logger.error("[index_build] native index_builder failed: %s", e)
        raise HTTPException(500, f"native index_builder failed: {e}")

    log_mem("[index_build] native: after index_builder")
    logger.info("[index_build] NATIVE done")



def _docx_to_text(raw: bytes) -> str:
    if Document is None:
        raise HTTPException(500, "python-docx not installed")
    f = io.BytesIO(raw)
    doc = Document(f)
    parts = [p.text for p in doc.paragraphs if p.text]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append(" ".join(c.text for c in row.cells if c.text))
    return "\n".join(parts)


def _guess_ext(fname: str) -> str:
    fname = (fname or "").lower()
    if fname.endswith((".txt", ".log", ".md")):
        return "txt"
    if fname.endswith(".docx"):
        return "docx"
    if fname.endswith(".pdf"):
        return "pdf"
    return ""


def _collect_corpus_stats() -> Tuple[int, Set[str]]:
    """Подсчёт строк и множества doc_id в corpus.jsonl."""
    lines = 0
    ids: Set[str] = set()
    if not CORPUS_JSONL.exists():
        return 0, set()

    with open(CORPUS_JSONL, "r", encoding="utf-8") as f:
        for line in f:
            s = line.strip()
            if not s:
                continue
            lines += 1
            try:
                obj = json.loads(s)
            except Exception:
                continue
            did = obj.get("doc_id")
            if did:
                ids.add(did)
    return lines, ids


def _collect_index_ids_safe() -> Set[str]:
    """doc_id из текущего index.json, если он есть и валиден."""
    if not INDEX_JSON.exists():
        return set()
    try:
        from ..services.search.index_search import load_index
        idx = load_index()
    except Exception:
        return set()
    docs_meta = idx.get("docs_meta") or {}
    return set(docs_meta.keys())


def _ingest_single(
    raw: bytes,
    fname: str,
    normalize: bool,
    save_docx: bool,
    ocr: bool,
    ocr_mode: str,
    title: str | None = None,
    author: str | None = None,
) -> Dict[str, Any]:
    """
    Общий пайплайн для одного файла.
    Возвращает dict с doc_id, токенами и пр.
    """
    log_mem(f"ingest_single: start fname={fname}")
    rt_cfg = get_runtime_cfg()
    ocr_lang = rt_cfg.ocr.lang
    ocr_workers = rt_cfg.ocr.workers

    ctype_guess = _guess_ext(fname)
    text: str

    logger.info(
        "[ingest_single] start fname=%s, ext=%s, bytes=%d, normalize=%s, save_docx=%s, ocr=%s, ocr_mode=%s",
        fname,
        ctype_guess,
        len(raw),
        normalize,
        save_docx,
        ocr,
        ocr_mode,
    )

    # 1) извлечение текста
    if ctype_guess == "txt":
        text = _bytes_to_text_utf8(raw)

    elif ctype_guess == "docx":
        try:
            text = extract_docx_text(raw)
        except RuntimeError as e:
            logger.error("[ingest_single] docx extract failed: %s", e)
            raise HTTPException(500, str(e))

    elif ctype_guess == "pdf":
        # heavy PDF normalize for indexing
        try:
            text = extract_and_normalize_pdf(
                raw,
                try_ocr=ocr,
                dpi=200,
                psm=3,
                return_debug=False,
                ocr_mode=ocr_mode,
                lang=ocr_lang,
                ocr_workers=ocr_workers,
            )
        except Exception as e:
            logger.error("[ingest_single] PDF extract failed: %s", e)
            raise HTTPException(500, f"PDF extract failed: {e}")

        # сохранить DOCX-версию
        if save_docx:
            try:
                docx_bytes = smart_pdf_to_docx(
                    raw,
                    try_ocr=ocr,
                    ocr_mode=ocr_mode,
                    force_ocr=False,
                    lang=ocr_lang,
                    ocr_workers=ocr_workers,
                )
                UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
                base = hashlib.sha256(raw).hexdigest()[:12]
                out_path = UPLOAD_DIR / f"{base}.docx"
                out_path.write_bytes(docx_bytes)
            except Exception as e:
                logger.error(f"docx conversion failed for {fname}: {e}")

    else:
        logger.error("[ingest_single] unsupported extension fname=%s", fname)
        raise HTTPException(400, f"unsupported file extension: {fname or 'unknown'}")

    # 2) нормализация
    if normalize:
        text = normalize_for_shingles(text)

    if not text.strip():
        logger.warning("[ingest_single] empty text after extraction fname=%s", fname)
        raise HTTPException(400, "empty text after extraction")

    if not title:
        base_name = Path(fname).name if fname else ""
        if "." in base_name:
            title = base_name.rsplit(".", 1)[0]
        else:
            title = base_name or ""

    if author is None:
        author = ""

    # 3) запись в корпус
    CORPUS_JSONL.parent.mkdir(parents=True, exist_ok=True)
    doc_id = f"doc_{hashlib.sha256(text.encode('utf-8')).hexdigest()[:8]}"

    with file_lock(CORPUS_JSONL.with_suffix(".lock")):
        with open(CORPUS_JSONL, "ab") as f:
            f.write(
                orjson.dumps(
                    {
                        "doc_id": doc_id,
                        "text": text,
                        "title": title,
                        "author": author,
                    }
                )
                + b"\n"
            )

    toks = simple_tokens(text)

    logger.info(
        "[ingest_single] done doc_id=%s, filename=%s, title=%s, bytes=%d, tokens=%d",
        doc_id,
        fname,
        title,
        len(raw),
        len(toks),
    )
    log_mem(f"ingest_single: done doc_id={doc_id}")

    return {
        "doc_id": doc_id,
        "filename": fname,
        "title": title,
        "author": author,
        "bytes": len(raw),
        "tokens": len(toks),
    }


def _iter_jsonl(path) -> Iterator[Dict[str, Any]]:
    with open(path, "r", encoding="utf-8") as f:
        for i, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                if "doc_id" in obj and "text" in obj:
                    obj["_line_no"] = i
                    yield obj
            except json.JSONDecodeError:
                continue


# ---------- routes ----------
@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    normalize: bool = Query(True),
    save_docx: bool = Query(True),
    ocr: bool = Query(True, description="Включить OCR, если PDF пустой"),
    ocr_mode: str = Query(
        "speed",
        pattern="^(speed|balanced|quality)$",
        description="OCR режим: speed / balanced / quality",
    ),
    title: str | None = Query(
        None,
        description="Название документа (если не указано — берётся из имени файла)",
    ),
    author: str | None = Query(
        None,
        description="Автор документа (если не указано — пустая строка)",
    ),
):
    fname = (file.filename or "").lower()
    raw = await file.read()

    logger.info(
        "[upload_file] start fname=%s, bytes=%d, normalize=%s, save_docx=%s, ocr=%s, ocr_mode=%s",
        fname,
        len(raw),
        normalize,
        save_docx,
        ocr,
        ocr_mode,
    )
    log_mem("upload_file: before _ingest_single")

    result = _ingest_single(
        raw=raw,
        fname=fname,
        normalize=normalize,
        save_docx=save_docx,
        ocr=ocr,
        ocr_mode=ocr_mode,
        title=title,
        author=author,
    )

    logger.info(
        "[upload_file] done doc_id=%s, tokens=%d",
        result.get("doc_id"),
        result.get("tokens"),
    )
    log_mem(f"upload_file: after _ingest_single doc_id={result.get('doc_id')}")

    return result


# ---------- массовая загрузка из ZIP ----------
@router.post("/upload/zip_admin")
async def upload_zip_admin(
    file: UploadFile = File(...),
    normalize: bool = Query(True),
    save_docx: bool = Query(True),
    ocr: bool = Query(True, description="Включить OCR для PDF внутри архива"),
    ocr_mode: str = Query(
        "speed",
        pattern="^(speed|balanced|quality)$",
        description="OCR режим: speed / balanced / quality",
    ),
    batch_size: int = Query(
        100,
        ge=1,
        le=1000,
        description="Сколько файлов обрабатывать в одном батче внутри архива",
    ),
):
    """
    Тяжёлый ZIP-роутер для себя: много файлов, может работать долго.
    Обрабатывает архив батчами по batch_size файлов, между батчами делает gc.collect().
    Логирует прогресс: N/total (~X.X%).
    """
    try:
        file.file.seek(0)
        zf = zipfile.ZipFile(file.file)
    except Exception as e:
        logger.error("[upload_zip_admin] not a valid ZIP: %s", e)
        raise HTTPException(400, f"not a valid ZIP: {e}")

    from starlette.concurrency import run_in_threadpool
    import gc as _gc

    # сначала отфильтруем валидные entries, чтобы знать total
    infos: List[zipfile.ZipInfo] = []
    for info in zf.infolist():
        name = info.filename.rsplit("/", 1)[-1]
        if info.is_dir():
            continue
        if not name:
            continue
        if name.startswith("__MACOSX"):
            continue
        infos.append(info)

    total = len(infos)
    if total == 0:
        logger.info(
            "[upload_zip_admin] empty archive=%s (no valid files)", file.filename
        )
        return {
            "archive_name": file.filename or "archive.zip",
            "total_files": 0,
            "processed": 0,
            "items": [],
        }

    logger.info(
        "[upload_zip_admin] start: archive=%s, files=%d, batch_size=%d, "
        "normalize=%s, save_docx=%s, ocr=%s, ocr_mode=%s",
        file.filename,
        total,
        batch_size,
        normalize,
        save_docx,
        ocr,
        ocr_mode,
    )
    log_mem("upload_zip_admin: start")

    items: List[Dict[str, Any]] = []
    processed = 0
    idx = 0  # глобальный счётчик по файлам

    for batch in _chunked(infos, batch_size):
        for info in batch:
            idx += 1
            name = info.filename.rsplit("/", 1)[-1]

            try:
                raw = zf.read(info)
            except Exception as e:
                status = "error"
                items.append(
                    {
                        "filename": name,
                        "status": status,
                        "error": f"cannot read from zip: {e}",
                    }
                )
                pct = idx * 100.0 / total
                logger.error(
                    "[upload_zip_admin] %d/%d (%.1f%%) filename=%s, status=%s, err=%s",
                    idx,
                    total,
                    pct,
                    name,
                    status,
                    e,
                )
                continue

            try:
                res = await run_in_threadpool(
                    _ingest_single,
                    raw,
                    name,
                    normalize,
                    save_docx,
                    ocr,
                    ocr_mode,
                    None,  # title
                    None,  # author
                )
                processed += 1
                res_out = dict(res)
                res_out.setdefault("status", "ok")
                items.append(res_out)
                status = "ok"
                pct = idx * 100.0 / total
                logger.info(
                    "[upload_zip_admin] %d/%d (%.1f%%) filename=%s, status=%s, doc_id=%s",
                    idx,
                    total,
                    pct,
                    name,
                    status,
                    res_out.get("doc_id"),
                )
            except HTTPException as e:
                status = "skipped"
                items.append(
                    {
                        "filename": name,
                        "status": status,
                        "error": e.detail,
                    }
                )
                pct = idx * 100.0 / total
                logger.warning(
                    "[upload_zip_admin] %d/%d (%.1f%%) filename=%s, status=%s, err=%s",
                    idx,
                    total,
                    pct,
                    name,
                    status,
                    e.detail,
                )
            except Exception as e:
                status = "error"
                items.append(
                    {
                        "filename": name,
                        "status": status,
                        "error": str(e),
                    }
                )
                pct = idx * 100.0 / total
                logger.error(
                    "[upload_zip_admin] %d/%d (%.1f%%) filename=%s, status=%s, err=%s",
                    idx,
                    total,
                    pct,
                    name,
                    status,
                    e,
                )

        # после каждого батча слегка чистим память
        _gc.collect()
        log_mem(
            f"upload_zip_admin: after batch, processed={processed}/{total}"
        )
        logger.info(
            "[upload_zip_admin] batch finished: processed=%d/%d",
            processed,
            total,
        )

    logger.info(
        "[upload_zip_admin] done: archive=%s, total=%d, processed=%d",
        file.filename,
        total,
        processed,
    )
    log_mem("upload_zip_admin: done")

    return {
        "archive_name": file.filename or "archive.zip",
        "total_files": total,
        "processed": processed,
        "items": items,
    }


@router.post("/build")
def build_index(
    mode: str = Query(
        "incremental",
        pattern="^(full|incremental)$",
        description="full — пересобрать всё, incremental — дозалить поверх текущего",
    ),
    max_new_docs_per_build: int | None = Query(
        None,
        ge=1,
        description=(
            "Максимум НОВЫХ документов за один прогон build. "
            "Используется как размер батча."
        ),
    ),
    process_all_pending: bool = Query(
        False,
        description=(
            "Если false — делаем ОДИН прогон build_index_json на max_new_docs_per_build.\n"
            "Если true (только для incremental) — крутим несколько прогонов подряд, "
            "каждый не больше max_new_docs_per_build, пока не закончатся "
            "неиндексированные документы или батч перестанет что-то добавлять."
        ),
    ),
):
    rt_cfg = get_runtime_cfg()

    logger.info(
        "[index_build] === STAGE 0: API /build called ===\n"
        "  mode=%s\n"
        "  max_new_docs_per_build=%s\n"
        "  process_all_pending=%s",
        mode,
        max_new_docs_per_build or "None",
        process_all_pending,
    )
    log_mem("index_build: STAGE 0 before stats")

    incremental = mode == "incremental"

    # для full смысла в process_all_pending нет — всегда один большой прогон
    if mode == "full" and process_all_pending:
        logger.warning(
            "[index_build] mode=full: process_all_pending=true не имеет смысла, "
            "будет выполнен один полный прогон без батчинга"
        )
        process_all_pending = False

    # ── STAGE 1: статы ДО ──────────────────────────────────────────────────────
    corpus_lines_before, corpus_ids = _collect_corpus_stats()
    index_ids_before = _collect_index_ids_safe()
    unindexed_before = corpus_ids - index_ids_before

    logger.info(
        "[index_build] STAGE 1/3: BEFORE\n"
        "  corpus.jsonl: lines=%d, doc_ids=%d\n"
        "  index.json:  indexed_docs=%d\n"
        "  unindexed_docs=%d",
        corpus_lines_before,
        len(corpus_ids),
        len(index_ids_before),
        len(unindexed_before),
    )
    log_mem("index_build: STAGE 1 BEFORE")

    # safety: если просим process_all_pending=true, но не задали размер батча — заставим задать
    if process_all_pending and max_new_docs_per_build is None and incremental:
        raise HTTPException(
            400,
            "process_all_pending=true требует задать max_new_docs_per_build "
            "для инкрементального режима",
        )

    # ── STAGE 2: BUILD (Python index.json) ─────────────────────────────────────
    logger.info(
        "[index_build] STAGE 2/3: BUILD start\n"
        "  incremental=%s, max_new_docs_per_build=%s, process_all_pending=%s",
        incremental,
        max_new_docs_per_build or "None",
        process_all_pending,
    )
    log_mem("index_build: STAGE 2 BUILD start")

    last_idx = None
    total_batches = 0

    def _run_single_build(prev_docs_count: int | None = None) -> tuple[dict, int]:
        """Запускает build_index_json и возвращает (idx, delta_new_docs)."""
        log_mem("index_build: before build_index_json")
        try:
            idx = build_index_json(
                cfg=rt_cfg.index.model_dump(),
                incremental=incremental,
                max_new_docs_per_build=max_new_docs_per_build,
            )
        except FileNotFoundError:
            logger.error("[index_build] corpus.jsonl not found")
            raise HTTPException(400, "corpus.jsonl not found")
        except ValueError as e:
            logger.error("[index_build] cannot build index: %s", e)
            raise HTTPException(400, f"cannot build index: {e}")

        log_mem("index_build: after build_index_json")

        docs_meta = idx.get("docs_meta") or {}
        cur_docs_count = len(docs_meta)
        if prev_docs_count is None:
            delta = 0
        else:
            delta = cur_docs_count - prev_docs_count

        logger.info(
            "[index_build] build pass done\n"
            "  docs_meta=%d\n"
            "  inverted.k9=%d\n"
            "  inverted.k13=%d\n"
            "  delta_new_docs=%d",
            cur_docs_count,
            len(idx["inverted_doc"].get("k9") or {}),
            len(idx["inverted_doc"].get("k13") or {}),
            delta,
        )
        log_mem("index_build: after build pass")
        return idx, delta

    # режим: либо один прогон, либо batched-multi-pass
    if not process_all_pending or not incremental:
        last_idx, _ = _run_single_build(prev_docs_count=len(index_ids_before))
        total_batches = 1
    else:
        prev_count = len(index_ids_before)
        max_loops = 10_000
        for i in range(1, max_loops + 1):
            logger.info(
                "[index_build] === incremental batch #%d START ===", i
            )
            idx, delta = _run_single_build(prev_docs_count=prev_count)
            last_idx = idx
            total_batches = i

            docs_meta = idx.get("docs_meta") or {}
            cur_count = len(docs_meta)
            prev_count = cur_count

            index_ids_now = set(docs_meta.keys())
            unindexed_now = corpus_ids - index_ids_now

            logger.info(
                "[index_build] incremental batch #%d AFTER\n"
                "  indexed_docs=%d\n"
                "  unindexed_docs=%d\n"
                "  delta_new_docs=%d",
                i,
                len(index_ids_now),
                len(unindexed_now),
                delta,
            )
            log_mem(f"index_build: after incremental batch #{i}")

            if delta <= 0 or not unindexed_now:
                logger.info(
                    "[index_build] incremental batches DONE at #%d: "
                    "delta=%d, unindexed_left=%d",
                    i,
                    delta,
                    len(unindexed_now),
                )
                break
        else:
            logger.warning(
                "[index_build] reached max_loops in incremental batched mode"
            )

    idx = last_idx or {}

    logger.info(
        "[index_build] STAGE 2/3: BUILD done\n"
        "  batches=%d\n"
        "  docs_meta=%d\n"
        "  inverted.k9=%d\n"
        "  inverted.k13=%d",
        total_batches,
        len(idx.get("docs_meta") or {}),
        len(idx.get("inverted_doc", {}).get("k9") or {}),
        len(idx.get("inverted_doc", {}).get("k13") or {}),
    )
    log_mem("index_build: STAGE 2 BUILD done")

    # ── STAGE 3: RELOAD Python-кэша ────────────────────────────────────────────
    logger.info("[index_build] STAGE 3/3: RELOAD index cache (python)")
    log_mem("index_build: STAGE 3 before clear_index_cache/load_index_cached")

    clear_index_cache()
    _ = load_index_cached()

    log_mem("index_build: after load_index_cached")

    index_ids_after = _collect_index_ids_safe()
    unindexed_after = corpus_ids - index_ids_after
    delta_indexed = len(index_ids_after) - len(index_ids_before)

    logger.info(
        "[index_build] STAGE 3/3: AFTER (python index)\n"
        "  corpus.jsonl: lines=%d, doc_ids=%d\n"
        "  index.json:  indexed_docs=%d\n"
        "  unindexed_docs=%d\n"
        "  delta_indexed=%d\n"
        "  batches=%d\n"
        "  process_all_pending=%s\n"
        "  max_new_docs_per_build=%s",
        corpus_lines_before,
        len(corpus_ids),
        len(index_ids_after),
        len(unindexed_after),
        delta_indexed,
        total_batches,
        process_all_pending,
        max_new_docs_per_build or "None",
    )
    log_mem("index_build: STAGE 3 AFTER")

    # ── STAGE 4: C++ native index ──────────────────────────────────────────────
    _build_native_index()

    logger.info(
        "[index_build] API /build finished: "
        "docs=%d, k5=%d, k9=%d, k13=%d, mode=%s, "
        "process_all_pending=%s, batches=%d, max_new_docs_per_build=%s",
        len(idx.get("docs_meta") or {}),
        len(idx.get("inverted_doc", {}).get("k5") or {}),
        len(idx.get("inverted_doc", {}).get("k9") or {}),
        len(idx.get("inverted_doc", {}).get("k13") or {}),
        mode,
        process_all_pending,
        total_batches,
        max_new_docs_per_build or "None",
    )
    log_mem("index_build: done")

    return {
        "index_path": str(INDEX_JSON),
        "mode": mode,
        "max_new_docs_per_build": max_new_docs_per_build,
        "process_all_pending": process_all_pending,
        "batches": total_batches,
        "stats_before": {
            "corpus_lines": corpus_lines_before,
            "corpus_docs": len(corpus_ids),
            "indexed_docs": len(index_ids_before),
            "unindexed_docs": len(unindexed_before),
        },
        "stats_after": {
            "corpus_lines": corpus_lines_before,
            "corpus_docs": len(corpus_ids),
            "indexed_docs": len(index_ids_after),
            "unindexed_docs": len(unindexed_after),
            "delta_indexed": delta_indexed,
        },
        "docs": len(idx.get("docs_meta") or {}),
        "k5": len(idx.get("inverted_doc", {}).get("k5") or {}),
        "k9": len(idx.get("inverted_doc", {}).get("k9") or {}),
        "k13": len(idx.get("inverted_doc", {}).get("k13") or {}),
    }
@router.delete(
    "/corpus/cleanup",
    summary="Удалить полностью базу или только неиндексированные документы",
)
def corpus_cleanup(
    delete_all: bool = Query(
        False,
        description="true — удалить ВСЮ базу и индекс; false — удалить только неиндексированные",
    ),
):
    # ===================== ВЕСЬ CORPUS + INDEX =====================
    if delete_all:
        log_mem("corpus_cleanup: delete_all")

        removed: list[str] = []
        for p in (CORPUS_JSONL, INDEX_JSON):
            if p.exists():
                try:
                    os.remove(p)
                    removed.append(str(p))
                except OSError as e:
                    raise HTTPException(500, f"cannot remove {p.name}: {e}")

        clear_index_cache()

        logger.info("[corpus_cleanup] delete_all removed=%s", removed)
        log_mem("corpus_cleanup: delete_all done")
        return {
            "status": "ok",
            "mode": "all",
            "removed": removed,
        }

    # ===================== УДАЛИТЬ ТОЛЬКО НЕИНДЕКСИРОВАННЫЕ =====================
    log_mem("corpus_cleanup: unindexed")

    if not CORPUS_JSONL.exists():
        raise HTTPException(404, "corpus.jsonl not found")

    try:
        idx = load_index()
    except FileNotFoundError:
        raise HTTPException(400, "index not found — сначала построй индекс")
    except Exception as e:
        raise HTTPException(500, f"error loading index: {e}")

    docs_meta = idx.get("docs_meta") or {}
    indexed_ids = set(docs_meta.keys())
    if not indexed_ids:
        raise HTTPException(400, "index is empty, nothing to clean")

    lock_path = CORPUS_JSONL.with_suffix(".lock")
    tmp_path = CORPUS_JSONL.with_suffix(".tmp")

    total_lines = 0
    kept_lines = 0
    removed_lines = 0
    removed_ids: list[str] = []

    with file_lock(lock_path):
        with open(CORPUS_JSONL, "rb") as src, open(tmp_path, "wb") as dst:
            for line in src:
                line_stripped = line.strip()
                if not line_stripped:
                    continue

                total_lines += 1

                try:
                    obj = orjson.loads(line_stripped)
                except Exception:
                    dst.write(line)
                    kept_lines += 1
                    continue

                did = obj.get("doc_id")
                if did and did not in indexed_ids:
                    removed_lines += 1
                    removed_ids.append(did)
                    continue

                dst.write(line)
                kept_lines += 1

        os.replace(tmp_path, CORPUS_JSONL)

    logger.info(
        "[corpus_cleanup] unindexed: total=%d kept=%d removed=%d",
        total_lines,
        kept_lines,
        removed_lines,
    )
    log_mem("corpus_cleanup: unindexed done")

    return {
        "status": "ok",
        "mode": "unindexed",
        "total": total_lines,
        "kept": kept_lines,
        "removed": removed_lines,
        "removed_doc_ids_sample": removed_ids[:50],
    }


# 2) ЧИСТАЯ СТАТИСТИКА
