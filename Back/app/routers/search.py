import io
from fastapi import APIRouter, HTTPException, UploadFile, File, Query, Body
from pydantic import BaseModel
from ..services.index_search import search_cached, load_index

# Пытаться подключить парсер DOCX
try:
    from docx import Document  # python-docx
except Exception:
    Document = None

router = APIRouter(prefix="/api", tags=["Search"])

class SearchReq(BaseModel):
    query: str
    top: int = 5

def _ensure_index():
    try:
        load_index()
    except FileNotFoundError:
        raise HTTPException(400, "index.json not found; build first")
    except ValueError:
        raise HTTPException(500, "invalid index structure")

def _docx_to_text(raw: bytes) -> str:
    if Document is None:
        raise HTTPException(500, "python-docx is not installed")
    f = io.BytesIO(raw)
    doc = Document(f)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append(" ".join(c.text for c in row.cells))
    text = "\n".join(p for p in parts if p)
    return text

@router.post("/search")
def search(req: SearchReq):
    _ensure_index()
    return search_cached(req.query.strip(), top=req.top)

@router.post("/search-raw")
async def search_raw(
    body: str = Body(..., media_type="text/plain", description="Сырой текст запроса"),
    top: int = Query(5, ge=1, le=50),
):
    _ensure_index()
    text = (body or "").strip()
    if not text:
        raise HTTPException(400, "empty body")
    return search_cached(text, top=top)

@router.post("/search-file")
async def search_file(
    file: UploadFile = File(...),
    top: int = Query(5, ge=1, le=50),
):
    """
    Поиск по содержимому файла. Поддерживает text/* (.txt/.log/.md) и Word (.docx).
    """
    _ensure_index()

    ctype = (file.content_type or "").lower()
    fname = (file.filename or "").lower()
    raw = await file.read()

    # text/*
    if ctype.startswith("text/") or fname.endswith((".txt", ".log", ".md")):
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("utf-8", "ignore")

    # .docx (OOXML)
    elif (
        ctype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or fname.endswith(".docx")
    ):
        text = _docx_to_text(raw)

    # .doc (старый MS Word) — не поддерживаем надёжно без внешних утилит
    elif ctype == "application/msword" or fname.endswith(".doc"):
        raise HTTPException(415, "legacy .doc is not supported; convert to .docx or .txt")

    else:
        raise HTTPException(400, f"unsupported file type: {ctype or fname or 'unknown'}")

    text = (text or "").strip()
    if not text:
        raise HTTPException(400, "empty file after extraction")

    return search_cached(text, top=top)
