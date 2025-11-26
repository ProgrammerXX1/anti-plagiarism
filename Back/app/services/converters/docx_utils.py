# app/services/converters/docx_utils.py
import io
from typing import List

try:
    from docx import Document  # python-docx
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


def _ensure_docx_available() -> None:
    if not _HAS_DOCX:
        raise RuntimeError("python-docx is not installed")


def extract_docx_text(raw: bytes, skip_empty_cells: bool = True) -> str:
    """
    Унифицированное извлечение текста из DOCX:
    - абзацы
    - таблицы
    """
    _ensure_docx_available()

    doc = Document(io.BytesIO(raw))
    parts: List[str] = []

    # абзацы
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # таблицы
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = []
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t or not skip_empty_cells:
                    cells.append(t)
            if cells:
                parts.append(" ".join(cells))

    return "\n".join(parts)
