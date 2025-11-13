# app/services/pdf_convert.py
import io, subprocess, tempfile, os
from pathlib import Path
from typing import Optional
from .pdf_heavy import pytess_ocr_pdf

def pdf_to_docx_bytes(raw_pdf: bytes, start: int = 0, end: Optional[int] = None) -> bytes:
    """Конвертирует 'текстовый' PDF в DOCX. Без OCR."""
    from pdf2docx import Converter  # требует pip install pdf2docx
    with tempfile.TemporaryDirectory() as td:
        pdf_path  = Path(td) / "in.pdf"
        docx_path = Path(td) / "out.docx"
        pdf_path.write_bytes(raw_pdf)
        cv = Converter(str(pdf_path))
        try:
            cv.convert(str(docx_path), start=start, end=end)
        finally:
            cv.close()
        return docx_path.read_bytes()

def ocr_pdf_bytes(raw_pdf: bytes, lang: str = "rus+eng") -> bytes:
    """OCR для сканов → searchable PDF. Нужен ocrmypdf + tesseract."""
    with tempfile.TemporaryDirectory() as td:
        inp = Path(td) / "in.pdf"
        out = Path(td) / "out.pdf"
        inp.write_bytes(raw_pdf)
        cmd = [
            "ocrmypdf",
            "--skip-text",      # или "--force-ocr", но не оба
            "-l", lang,
            str(inp),
            str(out),
        ]
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if proc.returncode != 0 or not out.exists():
            raise RuntimeError(
                f"OCR failed: rc={proc.returncode}, err={proc.stderr.decode('utf-8', 'ignore')}"
            )
        return out.read_bytes()

def _build_docx_from_text(text: str) -> bytes:
    """Собрать DOCX из plain text."""
    from docx import Document
    doc = Document()
    for para in text.split("\n\n"):
        p = para.strip()
        if p:
            doc.add_paragraph(p)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def smart_pdf_to_docx(
    raw_pdf: bytes,
    try_ocr: bool = True,
    lang: str = "kaz+rus+eng",
    ocr_workers: int = 16,
    ocr_mode: str = "speed",     # speed / balanced / quality
    force_ocr: bool = False,     # <<< ГЛАВНЫЙ ПАРАМЕТР ДЛЯ ТЕБЯ
) -> bytes:
    """
    Варианты работы:

    1) force_ocr=True:
       - игнорируем pdf2docx
       - сразу распознаём OCR и собираем DOCX из текста

    2) force_ocr=False (авторежим):
       - пробуем pdf2docx
       - если в DOCX мало текста (скан), прогоняем OCR и собираем текстовый DOCX
    """
    # Если OCR вообще не нужен — просто конвертируем как есть
    if not try_ocr and not force_ocr:
        return pdf_to_docx_bytes(raw_pdf)

    # ─────────────────────────────────────────────────────────────
    # Вариант 1: жёстко форсим OCR → только текстовый DOCX
    # ─────────────────────────────────────────────────────────────
    if force_ocr:
        txt, dbg = pytess_ocr_pdf(
            raw_pdf,
            lang=lang,
            ocr_mode=ocr_mode,
            workers=ocr_workers,
        )
        if txt.strip():
            return _build_docx_from_text(txt)

        # Если OCR внезапно ничего не дал — фолбэк через ocrmypdf или тупо pdf2docx
        try:
            searchable = ocr_pdf_bytes(raw_pdf, lang=lang)
            return pdf_to_docx_bytes(searchable)
        except Exception:
            # совсем крайний случай
            return pdf_to_docx_bytes(raw_pdf)

    # ─────────────────────────────────────────────────────────────
    # Вариант 2: умный режим — сначала pdf2docx, потом OCR, если скан
    # ─────────────────────────────────────────────────────────────
    docx = pdf_to_docx_bytes(raw_pdf)

    # Пытаемся понять, есть ли в docx реальный текст
    text_len = 0
    try:
        from docx import Document
        d = Document(io.BytesIO(docx))
        for p in d.paragraphs:
            text_len += len(p.text.strip() or "")
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text_len += len(cell.text.strip() or "")
    except Exception:
        # если python-docx не установлен/сломался — будем считать, что текста мало
        text_len = 0

    # Если текста достаточно — считаем, что это не скан, OCR не нужен
    if text_len >= 500 or not try_ocr:
        return docx

    # Иначе — скан, делаем OCR и собираем текстовый DOCX
    txt, dbg = pytess_ocr_pdf(
        raw_pdf,
        lang=lang,
        ocr_mode=ocr_mode,
        workers=ocr_workers,
    )

    if txt.strip():
        return _build_docx_from_text(txt)

    # если и тут пусто — фолбэк через ocrmypdf
    try:
        searchable = ocr_pdf_bytes(raw_pdf, lang=lang)
        return pdf_to_docx_bytes(searchable)
    except Exception:
        return docx  # в самом крайнем случае вернём исходную конвертацию
